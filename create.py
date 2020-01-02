import os
import shutil
from docx import Document
from docx.shared import Inches

document = Document()


def createFiles(url):
    files = os.listdir(url)
    newDir = url+"\\png"
    for fileName in files:
        file_name, extension_name = os.path.splitext(fileName)
        if extension_name == '.png':
            shutil.move(url+"\\"+str(fileName), newDir)
        else:
            continue

    os.chdir(newDir)
    picTure = os.listdir(newDir)
    doc = Document()
    for png in picTure:
        file_name, extension_name = os.path.splitext(png)
        if extension_name == '.png':
            doc.add_picture(png)
        else:
            continue

    doc.save("发票.docx")